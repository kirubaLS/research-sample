"""Builds the non-technical proposal document for the Meet digest automation.
House style matches docs/diagrams/build_docx.py — Calibri, same heading palette."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED, INK, SLATE, MUTED = 'B33224', '16222F', '3E5062', '6F8296'
RULE, BAND = 'D9DEE4', 'F2F4F7'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)   # A4 portrait
sec.left_margin = sec.right_margin = Inches(0.85)
sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
USABLE = 6.57

n = doc.styles['Normal']; n.font.name = 'Calibri'; n.font.size = Pt(10.5)
n.paragraph_format.space_after = Pt(7)
for lvl, (sz, col) in enumerate([(22, RED), (14.5, INK), (11.5, SLATE)], start=1):
    s = doc.styles[f'Heading {lvl}']
    s.font.size = Pt(sz); s.font.color.rgb = RGBColor.from_string(col)
    s.font.name = 'Calibri'; s.font.bold = True
    s.paragraph_format.space_before = Pt(16 if lvl < 3 else 11)
    s.paragraph_format.space_after = Pt(5)


def shade(cell, hexcol):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hexcol)
    cell._tc.get_or_add_tcPr().append(el)


def para(txt, size=10.5, bold=False, italic=False, col=INK, space=7, align=None, style=None):
    p = doc.add_paragraph(style=style); p.paragraph_format.space_after = Pt(space)
    if align: p.alignment = align
    r = p.add_run(txt); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(col)
    return p


def bullets(items, col=INK, size=10.5):
    for it in items:
        b = doc.add_paragraph(style='List Bullet')
        b.paragraph_format.space_after = Pt(3)
        bold_part, _, rest = it.partition('||')
        r = b.add_run(bold_part); r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(col)
        if rest:
            r.bold = True
            r2 = b.add_run(rest); r2.font.size = Pt(size); r2.font.color.rgb = RGBColor.from_string(col)


def callout(title, body, tone=RED):
    t = doc.add_table(rows=1, cols=1); t.autofit = False
    c = t.rows[0].cells[0]; c.width = Inches(USABLE)
    shade(c, BAND)
    c.paragraphs[0].text = ''
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title); r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(tone)
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body); r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def table(headers, rows, widths, sizes=9.8):
    t = doc.add_table(rows=1, cols=len(headers))
    # autofit=False is what emits <w:tblLayout w:type="fixed"/> in the schema-correct
    # position; adding that element by hand lands it after tblLook and breaks the file.
    t.style = 'Table Grid'; t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.width = Inches(widths[i]); shade(c, SLATE)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].text = ''
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string('FFFFFF')
    for row in rows:
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cells[i].paragraphs[0].text = ''
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
            r = cells[i].paragraphs[0].add_run(str(txt)); r.font.size = Pt(sizes)
            r.font.color.rgb = RGBColor.from_string(INK)
            if i == 0 and len(headers) > 2: r.bold = True
    return t

# ============================ COVER ============================
doc.add_heading("Automatic Follow-Up Notes for the Daily Stand-Up", level=1)
para("A proposal, and the step-by-step plan to set it up", 13, col=SLATE, space=16)

callout("In one minute",
        "Google Meet already records a written transcript of our daily stand-up. Nobody reads it — "
        "it is one long document of everything everyone said. This proposal turns that transcript "
        "into a short, personal email for each person, sent automatically within 15 minutes of the "
        "meeting ending. It uses tools we already pay for, nothing joins our meetings, and setting "
        "it up takes about half a day of one person's time.")

para("Prepared for management review · 31 August 2026", 9.5, col=MUTED, space=2)
para("Requires four decisions from you — see “What we need from you” on page 3.",
     9.5, italic=True, col=MUTED, space=14)

# ============================ PROBLEM ============================
doc.add_heading("1. The problem today", level=2)
para("Every stand-up produces a transcript. In practice it goes unread, because reading forty "
     "minutes of conversation to find the two lines that concern you is slower than just asking "
     "someone. Three things follow from that:", space=6)
bullets([
    "Commitments get lost.|| Someone agrees to do something, and by Thursday nobody can remember "
    "who owns it or what date was said.",
    "People miss things said about their work|| while they were half-listening, or on a day they "
    "were not in the room.",
    "The same questions get asked twice,|| because the answer was given in a meeting and never "
    "written down anywhere findable.",
])

doc.add_heading("2. What we are proposing", level=2)
para("Within 15 minutes of the stand-up ending, each person receives a short email containing only "
     "what concerns them. Not a summary of the meeting — a summary of their part in it.", space=6)
para("Importantly, this is not a transcript of what each person said. People already know what "
     "they said. The useful content is mostly what OTHER people said about them, which is exactly "
     "what is impossible to find by skimming.", space=10)

doc.add_heading("What one person would receive", level=3)

t = doc.add_table(rows=1, cols=1); c = t.rows[0].cells[0]; c.width = Inches(USABLE)
shade(c, 'FFFFFF')
c.paragraphs[0].text = ''
def mail(txt, size=10, bold=False, italic=False, col=INK, space=3, indent=0.0):
    p = c.add_paragraph() if c.paragraphs[0].runs or indent else c.paragraphs[0]
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(txt); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = RGBColor.from_string(col)

mail("Subject:  Your notes — 28 Aug, 9:30 am", 10, bold=True, col=SLATE, space=6)
mail("You agreed to two things and Ravi is waiting on one answer from you.", 10, space=8)
mail("YOUR COMMITMENTS", 9, bold=True, col=RED, space=3)
mail("•  Send the revised pricing sheet — due 29 August", 10, space=1, indent=0.12)
mail("“I'll get you the revised pricing sheet by tomorrow.”", 9, italic=True, col=MUTED, space=6, indent=0.25)
mail("QUESTIONS WAITING ON YOU", 9, bold=True, col=RED, space=3)
mail("•  Ravi asked whether the Chennai demo moves to Friday. You did not answer.", 10, space=6, indent=0.12)
mail("YOU WERE MENTIONED", 9, bold=True, col=RED, space=3)
mail("•  Priya: “Kingshuk has the client history on this, he should be on the call.”", 10, space=6, indent=0.12)
mail("DECISIONS AFFECTING YOUR WORK", 9, bold=True, col=RED, space=3)
mail("•  The onboarding rewrite is paused until the September numbers are in.", 10, space=2, indent=0.12)

para("An illustration, not a real meeting. Sections with nothing in them are left out, and a person "
     "with nothing relevant that day gets no email at all.",
     9, italic=True, col=MUTED, space=12)

doc.add_page_break()

# ============================ HOW IT WORKS ============================
doc.add_heading("3. How it works", level=2)
para("Four steps, all automatic. There is nothing for anyone to click, during the meeting or after.",
     space=8)
table(["", "What happens", "Who does it"],
      [["1", "The stand-up runs as normal. Google Meet writes its transcript, as it "
             "already does today.", "Nobody — automatic"],
       ["2", "A small program checks every 15 minutes whether a new stand-up has finished. "
             "It runs on Google's servers, not on anyone's laptop.", "Nobody — automatic"],
       ["3", "The transcript is read by Google's Gemini AI, which sorts the conversation "
             "into what matters for each named person.", "Nobody — automatic"],
       ["4", "Each person gets their own email, sent from our own Google account.",
             "Nobody — automatic"]],
      [0.35, 4.3, 1.92])
para("", space=6)

callout("Nothing joins our meetings",
        "This is the question most people ask first. No bot appears in the call, no recording "
        "assistant, no new app, and nobody sees anything different during the stand-up. The "
        "program simply reads the transcript Google has already written, afterwards.", tone=SLATE)

# ============================ DECISIONS ============================
doc.add_heading("4. What we need from you", level=2)
para("Four decisions. The first two need someone with Google Workspace administrator access; "
     "we cannot proceed without them.", space=8)

table(["Decision", "Why it is needed", "Effort"],
      [["Turn on automatic transcription\nfor our organisation",
        "Today someone has to remember to press “Start transcript” every morning. If they forget, "
        "there is no transcript and no emails that day. A setting in the Google admin console makes "
        "it start by itself.",
        "5 minutes,\nadministrator"],
       ["Approve recreating the daily\nstand-up calendar invite",
        "Google's automatic-transcription setting does not apply to recurring meetings created "
        "before it was switched on. Ours was. The invite has to be deleted and recreated once, or "
        "the setting silently does nothing.",
        "5 minutes,\nmeeting owner"],
       ["Approve a small Google Cloud\nbilling account",
        "Roughly ₹100–200 per month. This is not about the cost — see section 6. It is what keeps "
        "our meeting content out of Google's model training.",
        "15 minutes,\nfinance approval"],
       ["Approve telling the team,\nand a start date",
        "Everyone should hear about this at a stand-up before the first email arrives. An email "
        "quoting a colleague, arriving unannounced, feels like monitoring even when it is helpful.",
        "One stand-up"]],
      [1.75, 3.45, 1.37], sizes=9.5)

doc.add_page_break()

# ============================ PROCEDURE ============================
doc.add_heading("5. Step-by-step procedure", level=2)
para("Twelve steps in four stages. Stage A is the part that needs you; the rest is technical setup "
     "and does not need management involvement beyond the go-live decision in stage D.", space=9)

doc.add_heading("Stage A — Permissions and settings  (day 1, about 30 minutes)", level=3)
table(["Step", "What happens", "Who", "Time"],
      [["1", "Confirm our Google Workspace plan includes automatic transcription "
             "(Business Plus or higher). If we are on Business Standard, transcripts still work "
             "but someone must press the button each morning.", "Administrator", "5 min"],
       ["2", "Switch on automatic transcription in the Google admin console for our team.",
             "Administrator", "5 min"],
       ["3", "Delete and recreate the daily stand-up calendar invite, so the new setting applies "
             "to it. Same time, same people.", "Meeting owner", "5 min"],
       ["4", "Run one ordinary stand-up and confirm a transcript appears afterwards. "
             "Nothing else can be tested until this works.", "Meeting owner", "1 day"]],
      [0.42, 4.25, 1.35, 0.55], sizes=9.3)

doc.add_heading("Stage B — Technical setup  (day 2, about half a day)", level=3)
table(["Step", "What happens", "Who", "Time"],
      [["5", "Create a Google Cloud project and switch on the Meet service, so the program is "
             "allowed to read our own transcripts.", "Engineer", "20 min"],
       ["6", "Create the Gemini AI access key and attach the billing account approved in "
             "stage A.", "Engineer", "15 min"],
       ["7", "Install the program. It is a script that lives inside our own Google account — "
             "there is no server to buy and no software to license.", "Engineer", "30 min"],
       ["8", "Enter the settings: our stand-up's meeting code, and the list of team members "
             "and their email addresses.", "Engineer", "10 min"],
       ["9", "Run the built-in health check, which confirms every part is connected and prints "
             "who it recognised. Nothing is sent.", "Engineer", "10 min"]],
      [0.42, 4.25, 1.35, 0.55], sizes=9.3)

doc.add_heading("Stage C — Trial with nothing sent to the team  (week 1)", level=3)
table(["Step", "What happens", "Who", "Time"],
      [["10", "Switch the automation on, with every email redirected to one person instead of "
              "the team. It runs by itself from this point.", "Engineer", "5 min"],
       ["11", "For one week, that person reads all the emails the team would have received and "
              "reports whether they are accurate and useful. Wording is adjusted based on what "
              "they find.", "Reviewer", "1 week"]],
      [0.42, 4.25, 1.35, 0.55], sizes=9.3)

doc.add_heading("Stage D — Go live  (week 2 onwards)", level=3)
table(["Step", "What happens", "Who", "Time"],
      [["12", "Announce it at a stand-up, then switch the emails through to the team. "
              "From here it needs no attention.", "Manager", "1 stand-up"]],
      [0.42, 4.25, 1.35, 0.55], sizes=9.3)

para("", space=4)
callout("Recommended: one extra month before step 12",
        "Rather than going straight to the team, keep all the emails going to one reviewer for a "
        "month and have them forward the good ones by hand. It is slower, but nothing reaches the "
        "team unreviewed while we are still learning what the AI gets wrong on our meetings.",
        tone=SLATE)

doc.add_page_break()

# ============================ COST ============================
doc.add_heading("6. What it costs", level=2)
table(["Item", "Cost", "Note"],
      [["Google Meet transcripts", "₹0", "Already included in our Workspace plan"],
       ["The program itself", "₹0", "Google Apps Script — included with Workspace, no server needed"],
       ["Sending the emails", "₹0", "Sent from our own Google account, no mail service"],
       ["Gemini AI", "about ₹100–200\nper month",
        "One stand-up a day. There is a free tier, but we should not use it — see below"],
       ["Software licences", "₹0", "No vendor, no subscription, no per-seat fee"],
       ["Setup", "about half a day\nof one person", "One time"]],
      [1.85, 1.35, 3.37], sizes=9.5)

para("", space=4)
callout("Why we should pay rather than use the free option",
        "Google's free AI tier is genuinely free and would cover our usage. But on the free tier "
        "Google's terms allow it to use submitted content to improve its products, and allow human "
        "reviewers to read it. Our stand-up transcripts would fall under that. On the paid tier, "
        "Google commits not to train on our content. For roughly ₹150 a month, that is not a close "
        "call — it is the single most important line in this document.")

# ============================ RISK ============================
doc.add_heading("7. Risks, and how each is handled", level=2)
table(["What could go wrong", "How it is handled"],
      [["Someone's private or sensitive comment gets emailed around.",
        "Anyone can say “off the record” in the meeting. The program recognises the phrase and "
        "sends nothing at all for that meeting — the whole day is dropped, not just that line."],
       ["The wrong person receives someone else's notes.",
        "The program only emails people on a list we type in by hand. If it cannot match a name "
        "with certainty — two people called Ravi, for instance — it refuses to guess and sends "
        "to neither, then reports it."],
       ["The AI misunderstands something and puts a wrong commitment in writing.",
        "Every item quotes the exact words it came from, so the reader can judge for themselves. "
        "Stage C exists to measure how often this happens before anyone else sees an email."],
       ["Nobody starts the transcript, so nothing is sent.",
        "This is what stage A step 2 prevents. If it still happens, the daily log records it and "
        "no incorrect email goes out."],
       ["People feel monitored.",
        "It is announced before it starts; anyone can be removed from the list on request; and "
        "the emails contain nothing that was not said openly in a meeting everyone attended."],
       ["The person who set it up leaves.",
        "There is nothing to maintain day to day. The setup is documented step by step, and the "
        "whole thing can be switched off by one person in under a minute."]],
      [2.35, 4.22], sizes=9.5)

doc.add_page_break()

# ============================ TIMELINE ============================
doc.add_heading("8. Timeline", level=2)
table(["When", "What", "Decision point"],
      [["Day 1", "Admin settings; recreate the stand-up invite", "Your approval to begin"],
       ["Day 2", "Confirm a transcript appears after a real stand-up", "—"],
       ["Day 3", "Technical setup and health check", "—"],
       ["Week 1", "Runs automatically; all emails go to one reviewer", "—"],
       ["End of week 1", "Review of accuracy and usefulness", "Continue, adjust, or stop"],
       ["Weeks 2–5", "Reviewer forwards the good ones by hand (recommended)", "—"],
       ["Week 6", "Announce at stand-up; emails go directly to the team", "Your go-live approval"]],
      [1.15, 3.75, 1.67], sizes=9.5)

para("At any point up to the final step, stopping costs nothing — no contract, no licence, and "
     "nothing has reached the team.", 10, italic=True, col=MUTED, space=12)

# ============================ FAQ ============================
doc.add_heading("9. Questions you are likely to be asked", level=2)

faq = [
    ("Will something join or record our meetings?",
     "No. Nothing joins the call and no new recording is made. Google Meet already writes a "
     "transcript; this only reads it afterwards."),
    ("Does our meeting content leave the company?",
     "The transcript text is sent to Google's Gemini service to be sorted, and comes straight back. "
     "It already sits in Google Workspace, so it does not move to a new company — but it does move "
     "to a different Google service, which is why the paid tier in section 6 matters."),
    ("Can someone opt out?",
     "Yes. Removing a name from the list takes seconds and stops their emails. They still appear in "
     "the transcript, as they do today."),
    ("What if the AI writes something wrong?",
     "Every point quotes the words it came from, so it can be checked at a glance. Stage C is "
     "designed to find out how often this happens before the team is involved."),
    ("Is anyone tied to maintaining this?",
     "No. Once running, it needs no daily attention. The most likely thing to go wrong is a "
     "transcript not starting, which the log reports."),
    ("Can we stop it?",
     "Yes, immediately, by one person, in under a minute. There is no contract or licence."),
    ("Why not just use one of the meeting-notes products on the market?",
     "Those work, but each sends a bot into the call, which every external guest can see, and they "
     "charge per person. This uses what we already have and stays within our own Google account."),
]
for q, a in faq:
    para(q, 10.5, bold=True, col=SLATE, space=2)
    para(a, 10.5, space=9)

para("", space=2)
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6)
pb = OxmlElement('w:pBdr'); tp = OxmlElement('w:top')
tp.set(qn('w:val'), 'single'); tp.set(qn('w:sz'), '6'); tp.set(qn('w:color'), RULE)
pb.append(tp)
# pBdr must precede spacing/ind in pPr — schema order is enforced, and appending
# it after them is what silently breaks the file for LibreOffice and Word.
p._p.get_or_add_pPr().insert(0, pb)

para("The technical setup guide, written for whoever performs stages B and C, is a separate "
     "document. It does not need to be read to make the decisions in section 4.",
     9.5, italic=True, col=MUTED, space=2)

out = "/home/user/research-sample/docs/Meeting Follow-Up Automation — Proposal and Plan.docx"
doc.save(out)

# python-docx ships <w:zoom w:val="bestFit"/>, but the schema requires w:percent.
# Repair it in place so the file validates and opens everywhere.
import zipfile, shutil, os
tmp = out + ".tmp"
with zipfile.ZipFile(out) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/settings.xml":
            data = data.replace(b'<w:zoom w:val="bestFit"/>', b'<w:zoom w:percent="100"/>')
        zout.writestr(item, data)
shutil.move(tmp, out)
print("saved:", out)
