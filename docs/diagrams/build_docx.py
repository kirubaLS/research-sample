from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)
for a in ('left_margin','right_margin'): setattr(sec,a,Inches(0.5))
sec.top_margin = Inches(0.55); sec.bottom_margin = Inches(0.5)
USABLE = 10.69

n = doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(10.5)
for lvl,(sz,col) in enumerate([(24,'B33224'),(16,'16222F'),(12.5,'3E5062')],start=1):
    s=doc.styles[f'Heading {lvl}']; s.font.size=Pt(sz); s.font.color.rgb=RGBColor.from_string(col); s.font.name='Calibri'; s.font.bold=True

def para(txt, size=10.5, bold=False, italic=False, col='16222F', space=6, align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(space)
    if align: p.alignment=align
    r=p.add_run(txt); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    r.font.color.rgb=RGBColor.from_string(col)
    return p

def figure(png, num, title, whatit, notes):
    doc.add_heading(f"Figure {num} — {title}", level=2)
    para(whatit, 11, italic=True, col='3E5062', space=8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(8)
    p.add_run().add_picture(png, width=Inches(USABLE))
    para("Notes", 10.5, bold=True, col='B33224', space=3)
    for nt in notes:
        b=doc.add_paragraph(style='List Bullet'); b.paragraph_format.space_after=Pt(2)
        r=b.add_run(nt); r.font.size=Pt(10)
        r.font.color.rgb=RGBColor.from_string('3E5062')
    doc.add_page_break()

# ---- cover
doc.add_heading("Project Yaadhum — Workflow Diagrams", level=1)
para("End-to-end view: frontend, backend, storage, methods, models and algorithms",
     13, col='3E5062', space=14)
para("CBSE Class X · NCERT · Tamil Nadu · Bharath International Sr. Sec. School, Krishnagiri",
     10.5, col='6F8296', space=4)
para("Six diagrams. Read them in order — each one zooms in on the previous.", 10.5, col='6F8296', space=16)

rows = [("Figure 1","End-to-end system flow","Both use cases, five stages, and every write to storage."),
        ("Figure 2","Frontend view","Three journeys — student, teacher, principal — and the role boundary."),
        ("Figure 3","Backend view","Seven layers, the method and library for each, and the single model."),
        ("Figure 4","Data storage map","Four stores, two boundaries, and the training corpus."),
        ("Figure 5","The core algorithm","Photograph to verified mark, in four steps."),
        ("Figure 6","Accuracy and efficiency","Eight defences and the cost cascade.")]
t = doc.add_table(rows=0, cols=3); t.style='Light Grid Accent 1'
for a,b,c in rows:
    cells=t.add_row().cells
    for i,txt in enumerate((a,b,c)):
        cells[i].text=''
        r=cells[i].paragraphs[0].add_run(txt); r.font.size=Pt(10)
        if i<2: r.bold=True
doc.add_page_break()

figure("d1.png", 1, "End-to-End System Flow",
 "The whole system on one page: what the user touches, what the server does, and what is written down at each stage.",
 ["Read left to right. Five stages — capture in the browser, API and queue, processing, verification, delivery — with the psychometric test on the top lane and the marks engine on the bottom lane.",
  "The two lanes share one backend, one student table and one report layer. They are two products in one application, not two applications.",
  "Dashed arrows are writes to storage. Every stage writes down; nothing is held only in memory.",
  "Only the box labelled '7-layer extraction' and the question-paper box contain a model. Scoring, validation, reconciliation and reporting are ordinary deterministic code.",
  "The question paper is processed once per paper and its Q-matrix is frozen. That one-time cost is amortised across all forty students, which is why that stage can afford two independent passes of the best available model."])

figure("d2.png", 2, "Frontend View — Three Journeys, One Application",
 "What each of the three users sees, in what order, and where the hard boundary between them sits.",
 ["Journey A is deliberately a dead end. The student's last screen is 'thank you' — no score, no Holland code, no stream suggestion. An interest result handed to a fifteen-year-old without a counsellor beside them is how this product would do harm.",
  "That boundary is a row-level security rule in PostgreSQL, not a hidden button. A bug in the API cannot leak a result to a student session.",
  "Journey B is the scanner. Note the retake loop: pressing retake on page 12 of 15 re-shoots page 12 and keeps its position. A scanner that makes a teacher restart the script is a scanner they abandon after the third student.",
  "The quality gate runs before the shutter is usable — blur, glare, coverage and skew must all pass. A bad photo caught here costs five seconds; caught later it costs a wrong report.",
  "Pages live only in the browser until Complete is pressed. The whole capture flow works with no network, and a 24-hour purge clears any shared staffroom laptop automatically.",
  "Journey C ends on an accuracy panel showing the live audit-sample error rate. Showing a principal your own measured error rate is what separates a defensible product from a demo."])

figure("d3.png", 3, "Backend View — Seven Layers, One Model",
 "The method, algorithm, library and output contract for every layer of the marks engine.",
 ["The organising rule: exactly one layer contains a model. L4 is the only place a neural network appears.",
  "The six other layers are OpenCV, SciPy and plain Python — deterministic, CPU-cheap, unit-testable against fixed images, and byte-identical on every run. They cannot hallucinate.",
  "This is what makes the open-source and paid plans interchangeable: swapping the recogniser is a configuration value, not a rewrite. Roughly 70% of the code is shared between the two plans.",
  "It is also why failures are diagnosable. When something is wrong, it is wrong in a named layer with its own test suite — not 'the AI got it wrong'.",
  "L2 is where your red-pen answer pays off, and L6 is where arithmetic replaces guesswork. Those two layers do more for accuracy than any model choice."])

figure("d4.png", 4, "Data Storage Map — What Lives Where, and What Crosses",
 "Four stores, two hard boundaries, and the corpus that makes 'we will train our own model later' actually possible.",
 ["Nothing is ever updated in place. A corrected mark is a new row in mark_event; a view resolves the current value by precedence. This gives you audit trail, OCR-versus-teacher reconciliation and correction history for free.",
  "Student identity lives in a restricted schema with separate credentials. The analytics and training layers only ever see a pseudonymous UUID.",
  "The training boundary is crossed by crops and labels only. Full page images — which contain the student's handwriting — never enter the training corpus.",
  "A mark crop is a picture of one handwritten digit. No name, no PII, no substantive handwriting. That is what makes pooling across schools defensible under the DPDP Act.",
  "The five rules at the foot of the diagram are the ones most teams break. The costliest is storing only the crops a human reviewed — that leaves you with a training set made entirely of hard cases, and no way to measure accuracy on easy ones.",
  "Volume: forty students × thirty questions × two assessments is 2,400 labelled crops per class per term. Ten classes gets you to roughly 24,000 in one term, which is comfortably past what a fine-tuned digit recogniser needs."])

figure("d5.png", 5, "The Core Algorithm — From Photograph to a Verified Mark",
 "The four steps that turn a red number written anywhere on a page into a mark bound to the right question.",
 ["Step 1 is the highest-leverage step in the entire system, and it exists only because you confirmed teachers mark in red and students write in black or blue. It turns one hard problem into two easy ones: question numbers are found in the student image, marks in the teacher image.",
  "The student's own arithmetic working — the single largest source of false positives in a generic OCR pipeline — is simply not present in the picture the mark detector looks at.",
  "Step 2's closed-vocabulary filter is the anti-hallucination control at the localisation layer: a question label is accepted only if it exists in the frozen Q-matrix, so an invented 'Q47' on a thirty-question paper disappears by construction, at zero cost.",
  "Step 3 is solved as a constrained assignment problem with the Hungarian algorithm, not with a pile of if-statements. Then the second pass fits the teacher's own layout convention from the confident bindings and re-solves — that is what turns coin-flips into correct answers, and it is the piece a competitor on a generic OCR API will not have.",
  "Step 4 is the oracle. A handwritten 1 and 3 are the most confusable pair on an Indian answer script; per-crop a model may prefer 3, but the arithmetic says otherwise and the arithmetic is ground truth.",
  "If no assignment satisfies every equation, the script is flagged rather than guessed — a missing page, an unmarked question or the teacher's own addition error all surface here instead of silently corrupting a report.",
  "Practical ask: if no total is written anywhere on the script, ask the school to add one. It is the cheapest accuracy improvement available anywhere in this design."])

figure("d6.png", 6, "How Accuracy and Efficiency Are Actually Achieved",
 "The eight defences that stop an error reaching a report, and the cost cascade that keeps the system cheap.",
 ["No pipeline containing a recogniser is literally 100% correct. What this design guarantees is different and stronger: every mark reaching a report is either arithmetically verified against a number the teacher wrote, or was confirmed by a human.",
  "Each of the eight defences catches something the others cannot. Removing any one of them opens a specific failure mode, which is why they are worth building even though they overlap.",
  "The cost cascade pays by stakes rather than by volume: the best model reads the question paper once, a cheap model reads the 1,200 mark crops, and disagreements escalate. Roughly Rs 150–350 per class per assessment.",
  "Overnight batch processing halves that again, and answer scripts are not latency-sensitive — scan during the day, extract at night.",
  "The audit sample at the foot is what converts every claim above into a fact. Ten per cent of scripts re-keyed by hand, including cells the system was confident about, with the error rate published on the principal's dashboard.",
  "Without it you have a system that feels accurate. With it you have one that can prove it — and if accuracy ever drifts, you learn it from your own dashboard rather than from a parent."])

# closing
doc.add_heading("What these diagrams commit us to", level=2)
for t_,b_ in [("One application, not two",
   "The interest test and the marks engine share auth, the student table, the report layer and the deployment. Adding the second never means rebuilding the first."),
  ("One model, six deterministic layers",
   "The recogniser is the only replaceable component. This is what lets you launch on a paid model now and migrate to a free one later by changing a configuration value."),
  ("Colour separation before anything else",
   "Red teacher ink versus black or blue student ink is the assumption the whole localisation strategy rests on. If a school uses a different marking pen, the system fits new hue centroids automatically from three unlabelled pages — but a red pen genuinely improves their results, and it is worth telling them so."),
  ("Arithmetic as the oracle",
   "Totals are what make the extraction verifiable rather than merely probable. Every additional total — grand, section, page — is another equation, and each one measurably improves accuracy."),
  ("Store everything from the first page captured",
   "Every prediction with its full distribution, every human correction, every disagreement, with consent recorded at capture time. This is not optional groundwork; it is the entire bridge from the paid pilot to the free production system.")]:
    para(t_, 11.5, bold=True, col='16222F', space=2)
    para(b_, 10.5, col='3E5062', space=10)

doc.add_heading("Still outstanding", level=2)
for s_ in ["Five answer scripts photographed page by page on a teacher's own phone, in the room where scanning will happen. Not flatbed scans — they look nothing like production. This calibrates the association cost weights in Figure 5 and the ink profile in Figure 3.",
           "Whether a total is written anywhere on the script — grand, per-section, per-page, or none. See the note on Figure 5.",
           "Your company logo, as SVG plus a PNG fallback, with brand colours. It appears on the first screen in Figure 2 and on every exported PDF.",
           "Training-consent wording for the school agreement, so consent_class in Figure 4 is set correctly from the very first page captured."]:
    b=doc.add_paragraph(style='List Number'); r=b.add_run(s_); r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string('3E5062')

out="/home/user/research-sample/docs/Yaadhum Workflow Diagrams.docx"
doc.save(out); print("saved", out)
